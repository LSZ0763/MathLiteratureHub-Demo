import json
import os
import re
from datetime import datetime, date as dt_date
from typing import List
from sqlalchemy.orm import Session
from lxml import etree
import latex2mathml.converter
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from app import models, schemas

MML_NS = "http://www.w3.org/1998/Math/MathML"
OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _omml(tag):
    return "{%s}%s" % (OMML_NS, tag)


def _mml(tag):
    return "{%s}%s" % (MML_NS, tag)


def _convert_mml_element(elem):
    """Recursively convert MathML element to OMML element(s)."""
    tag = etree.QName(elem).localname

    if tag in ("mi", "mn", "mo", "mtext"):
        r = etree.Element(_omml("r"))
        t = etree.SubElement(r, _omml("t"))
        t.text = elem.text or ""
        return [r]

    if tag == "mrow":
        result = []
        for child in elem:
            result.extend(_convert_mml_element(child))
        return result

    if tag == "msup":
        children = list(elem)
        if len(children) >= 2:
            sSup = etree.Element(_omml("sSup"))
            e = etree.SubElement(sSup, _omml("e"))
            for c in _convert_mml_element(children[0]):
                e.append(c)
            sup = etree.SubElement(sSup, _omml("sup"))
            for c in _convert_mml_element(children[1]):
                sup.append(c)
            return [sSup]
        return []

    if tag == "msub":
        children = list(elem)
        if len(children) >= 2:
            sSub = etree.Element(_omml("sSub"))
            e = etree.SubElement(sSub, _omml("e"))
            for c in _convert_mml_element(children[0]):
                e.append(c)
            sub = etree.SubElement(sSub, _omml("sub"))
            for c in _convert_mml_element(children[1]):
                sub.append(c)
            return [sSub]
        return []

    if tag == "msubsup":
        children = list(elem)
        if len(children) >= 3:
            sSubSup = etree.Element(_omml("sSubSup"))
            e = etree.SubElement(sSubSup, _omml("e"))
            for c in _convert_mml_element(children[0]):
                e.append(c)
            sub = etree.SubElement(sSubSup, _omml("sub"))
            for c in _convert_mml_element(children[1]):
                sub.append(c)
            sup = etree.SubElement(sSubSup, _omml("sup"))
            for c in _convert_mml_element(children[2]):
                sup.append(c)
            return [sSubSup]
        return []

    if tag == "mfrac":
        children = list(elem)
        if len(children) >= 2:
            f = etree.Element(_omml("f"))
            num = etree.SubElement(f, _omml("num"))
            for c in _convert_mml_element(children[0]):
                num.append(c)
            den = etree.SubElement(f, _omml("den"))
            for c in _convert_mml_element(children[1]):
                den.append(c)
            return [f]
        return []

    if tag == "msqrt":
        rad = etree.Element(_omml("rad"))
        deg = etree.SubElement(rad, _omml("deg"))
        e = etree.SubElement(rad, _omml("e"))
        for child in elem:
            for c in _convert_mml_element(child):
                e.append(c)
        return [rad]

    if tag == "mroot":
        children = list(elem)
        if len(children) >= 2:
            rad = etree.Element(_omml("rad"))
            deg = etree.SubElement(rad, _omml("deg"))
            for c in _convert_mml_element(children[1]):
                deg.append(c)
            e = etree.SubElement(rad, _omml("e"))
            for c in _convert_mml_element(children[0]):
                e.append(c)
            return [rad]
        return []

    if tag == "mfenced":
        d = etree.Element(_omml("d"))
        dPr = etree.SubElement(d, _omml("dPr"))
        open_str = elem.get("open", "(")
        close_str = elem.get("close", ")")
        if open_str:
            open_el = etree.SubElement(dPr, _omml("open"))
            open_el.text = open_str
        if close_str:
            close_el = etree.SubElement(dPr, _omml("close"))
            close_el.text = close_str
        e = etree.SubElement(d, _omml("e"))
        for child in elem:
            for c in _convert_mml_element(child):
                e.append(c)
        return [d]

    if tag == "mover":
        children = list(elem)
        if len(children) >= 2:
            acc = etree.Element(_omml("acc"))
            accPr = etree.SubElement(acc, _omml("accPr"))
            chr_text = "".join(children[1].itertext())
            chr_el = etree.SubElement(accPr, _omml("chr"))
            chr_el.text = chr_text
            e = etree.SubElement(acc, _omml("e"))
            for c in _convert_mml_element(children[0]):
                e.append(c)
            return [acc]
        return []

    if tag == "munder":
        children = list(elem)
        if len(children) >= 2:
            limLow = etree.Element(_omml("limLow"))
            e = etree.SubElement(limLow, _omml("e"))
            for c in _convert_mml_element(children[0]):
                e.append(c)
            lim = etree.SubElement(limLow, _omml("lim"))
            for c in _convert_mml_element(children[1]):
                lim.append(c)
            return [limLow]
        return []

    if tag == "munderover":
        children = list(elem)
        if len(children) >= 3:
            limUpp = etree.Element(_omml("limUpp"))
            e = etree.SubElement(limUpp, _omml("e"))
            for c in _convert_mml_element(children[0]):
                e.append(c)
            lim = etree.SubElement(limUpp, _omml("lim"))
            for c in _convert_mml_element(children[2]):
                lim.append(c)
            return [limUpp]
        return []

    if tag == "math":
        omath = etree.Element(_omml("oMath"))
        for child in elem:
            for c in _convert_mml_element(child):
                omath.append(c)
        return [omath]

    # fallback: extract text
    text = "".join(elem.itertext())
    if text:
        r = etree.Element(_omml("r"))
        t = etree.SubElement(r, _omml("t"))
        t.text = text
        return [r]
    return []


def _set_math_font(run, font_name="Cambria Math", size=11):
    """Set font for math-supporting text with CJK fallback."""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    font.size = Pt(size)


def _add_text_with_latex_to_para(para, text, size=11, bold=False, italic=False):
    """Add text to paragraph, converting inline LaTeX ($...$ or \\(...\\)) to OMML."""
    pattern = re.compile(r"\$([^$\n]+?)\$|\\\(([^\)\n]+?)\\\)")
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            run = para.add_run(text[last_end:m.start()])
            _set_math_font(run, "Cambria Math", size)
            run.bold = bold
            run.italic = italic
        latex_str = m.group(1) if m.group(1) is not None else m.group(2)
        math_run = para.add_run()
        _set_math_font(math_run, "Cambria Math", size)
        math_run.bold = bold
        math_run.italic = italic
        try:
            mml_xml = latex2mathml.converter.convert(latex_str)
            mml_root = etree.fromstring(mml_xml.encode())
            omml_elements = _convert_mml_element(mml_root)
            for omml in omml_elements:
                math_run._element.append(omml)
        except Exception:
            math_run.text = m.group(0)
        last_end = m.end()
    if last_end < len(text):
        run = para.add_run(text[last_end:])
        _set_math_font(run, "Cambria Math", size)
        run.bold = bold
        run.italic = italic


def _add_heading_math(doc, text, level=1):
    """Add heading with math font support and LaTeX formula conversion."""
    heading = doc.add_heading(level=level)
    # Remove default run added by python-docx
    for run in heading.runs:
        run._element.getparent().remove(run._element)
    size = 16 if level == 1 else 14 if level == 2 else 12
    _add_text_with_latex_to_para(heading, text, size=size, bold=True)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if level > 0 else WD_PARAGRAPH_ALIGNMENT.CENTER
    return heading


def _add_paragraph_math(doc, text, bold=False, italic=False):
    """Add paragraph with math font support and LaTeX formula conversion."""
    p = doc.add_paragraph()
    _add_text_with_latex_to_para(p, text, size=11, bold=bold, italic=italic)
    return p


def generate_docx(req: schemas.GenerateBriefingRequest, db: Session) -> models.Briefing:
    # Fetch papers with sorting
    query = db.query(models.Paper).filter(models.Paper.id.in_(req.paper_ids))
    papers = query.all()

    if not papers:
        raise ValueError("No papers found for briefing")

    # Apply sorting
    if req.sort_by:
        reverse = req.sort_order == "desc"
        if req.sort_by == "title":
            papers.sort(key=lambda p: (p.title or "").lower(), reverse=reverse)
        elif req.sort_by == "author":
            papers.sort(key=lambda p: (p.authors or "").lower(), reverse=reverse)
        elif req.sort_by == "citation_count":
            papers.sort(key=lambda p: p.citation_count or 0, reverse=reverse)
        elif req.sort_by == "date":
            papers.sort(key=lambda p: p.published_date or dt_date.min, reverse=reverse)
        elif req.sort_by == "journal":
            papers.sort(key=lambda p: (p.journal or "").lower(), reverse=reverse)

    # Build doc
    doc = Document()

    # Set default font for document
    style = doc.styles["Normal"]
    style.font.name = "Cambria Math"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")

    # Title
    title = req.title or f"文献综述简报 - {datetime.now().strftime('%Y-%m-%d')}"
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in heading.runs:
        _set_math_font(run, "Cambria Math", 20)

    doc.add_paragraph()

    # Table of Contents
    toc_heading = _add_heading_math(doc, "目录", level=1)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for idx, paper in enumerate(papers, 1):
        toc_text = f"{idx}. {paper.title or '(无标题)'}"
        p = doc.add_paragraph()
        _add_text_with_latex_to_para(p, toc_text, size=11)

    doc.add_page_break()

    # Body
    for idx, paper in enumerate(papers, 1):
        # Paper heading
        paper_title = f"{idx}. {paper.title or '(无标题)'}"
        _add_heading_math(doc, paper_title, level=2)

        # Authors
        authors = ""
        try:
            authors = ", ".join(json.loads(paper.authors or "[]"))
        except Exception:
            authors = paper.authors or ""
        _add_paragraph_math(doc, f"作者：{authors}", bold=True)

        # Journal/Source
        if paper.journal:
            _add_paragraph_math(doc, f"杂志：{paper.journal}")
        elif paper.source:
            _add_paragraph_math(doc, f"来源：{paper.source}")

        # Date
        if paper.published_date:
            _add_paragraph_math(doc, f"发表日期：{paper.published_date}")

        # DOI
        if paper.doi:
            _add_paragraph_math(doc, f"DOI：{paper.doi}")

        # Link
        if paper.pdf_link:
            _add_paragraph_math(doc, f"链接：{paper.pdf_link}")

        # Citation / Reference counts
        meta_text = f"参考文献数量：{paper.reference_count or 0}  |  被引用数量：{paper.citation_count or 0}"
        _add_paragraph_math(doc, meta_text)

        # Keywords
        if paper.keywords:
            try:
                kw_list = json.loads(paper.keywords)
                if kw_list:
                    _add_paragraph_math(doc, f"关键词：{', '.join(kw_list)}")
            except Exception:
                pass

        # Abstract
        _add_heading_math(doc, "摘要", level=3)
        _add_paragraph_math(doc, paper.summary_raw or "（暂无摘要）")

        doc.add_paragraph()

    # Save
    output_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "..", "data", "briefings"))
    os.makedirs(output_dir, exist_ok=True)
    filename = f"briefing_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    briefing = models.Briefing(
        title=title,
        format="docx",
        file_path=filepath,
        paper_ids=json.dumps(req.paper_ids),
    )
    db.add(briefing)
    db.commit()
    db.refresh(briefing)
    return briefing
